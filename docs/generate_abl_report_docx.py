from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path
import shutil
import subprocess
import tempfile
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


@dataclass(frozen=True)
class ReportMeta:
    title: str
    subtitle: str
    project_name: str
    report_date: str
    student_name: str
    usn_or_id: str
    class_and_section: str
    department: str
    institution: str
    guide_name: str


def _set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Ensure East Asia font matches too (Word sometimes overrides)
    rpr = style.element.rPr
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), "Times New Roman")


def _add_title_page(doc: Document, meta: ReportMeta) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta.title)
    run.bold = True
    run.font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta.subtitle)
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()  # spacer

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(meta.project_name)
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()  # spacer

    for label, value in [
        ("Date", meta.report_date),
        ("Student Name", meta.student_name),
        ("USN / ID", meta.usn_or_id),
        ("Class / Section", meta.class_and_section),
        ("Department", meta.department),
        ("Institution", meta.institution),
        ("Guide / Mentor", meta.guide_name),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        p.add_run(value)

    doc.add_page_break()


def _add_section_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"


def _iter_md_lines(md_text: str) -> Iterable[str]:
    for line in md_text.splitlines():
        yield line.rstrip("\n")


def _render_markdown_like(doc: Document, md_text: str) -> None:
    """A tiny Markdown-ish renderer.

    Supports:
    - #, ##, ### headings
    - bullet lists starting with '-' or '*'
    - blank lines
    - everything else as paragraph text

    This is intentionally simple to keep the output predictable.
    """

    in_list = False

    def end_list() -> None:
        nonlocal in_list
        in_list = False

    for raw in _iter_md_lines(md_text):
        line = raw.strip()

        if not line:
            end_list()
            doc.add_paragraph()
            continue

        if line.startswith("### "):
            end_list()
            _add_section_heading(doc, line[4:].strip(), level=3)
            continue

        if line.startswith("## "):
            end_list()
            _add_section_heading(doc, line[3:].strip(), level=2)
            continue

        if line.startswith("# "):
            end_list()
            _add_section_heading(doc, line[2:].strip(), level=1)
            continue

        if line.startswith("- ") or line.startswith("* "):
            in_list = True
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:].strip())
            continue

        if in_list and (raw.startswith("  - ") or raw.startswith("  * ")):
            p = doc.add_paragraph(style="List Bullet 2")
            p.add_run(line[2:].strip())
            continue

        end_list()
        doc.add_paragraph(line)


def build_docx(
    project_report_md_path: Path,
    out_docx_path: Path,
    meta: ReportMeta,
) -> None:
    md_text = project_report_md_path.read_text(encoding="utf-8")

    doc = Document()
    _set_document_defaults(doc)

    _add_title_page(doc, meta)

    _add_section_heading(doc, "Index", level=1)
    doc.add_paragraph("1. Overview")
    doc.add_paragraph("2. Technology Stack")
    doc.add_paragraph("3. Repository Structure")
    doc.add_paragraph("4. Functional Modules")
    doc.add_paragraph("5. Data Storage")
    doc.add_paragraph("6. Doctor Status Rules")
    doc.add_paragraph("7. Input Validation")
    doc.add_paragraph("8. Configuration & Runtime")
    doc.add_paragraph("9. Security Notes")
    doc.add_paragraph("10. Development Tasks & Fixes Applied")
    doc.add_paragraph("11. Known Limitations & Suggested Enhancements")
    doc.add_paragraph("12. How to Run")
    doc.add_page_break()

    _add_section_heading(doc, "Activity Report", level=1)
    doc.add_paragraph(
        "This document summarizes the Activity Based Learning (ABL) project work carried out for the "
        "Hospital OPD Management System."
    )
    doc.add_paragraph()

    _render_markdown_like(doc, md_text)

    doc.add_page_break()
    _add_section_heading(doc, "Conclusion", level=1)
    doc.add_paragraph(
        "The Hospital OPD Management System demonstrates a complete full-stack workflow for OPD operations "
        "including patient registration, doctor management, and appointment booking. The solution provides "
        "a clean UI and a REST API backed by MongoDB, with clear scope for future hardening via server-side JWT enforcement."
    )

    out_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx_path)


def _convert_docx_to_pdf(docx_path: Path, out_pdf_path: Path) -> None:
    """Convert DOCX to PDF.

    Tries (in order):
    1) Microsoft Word automation via pywin32 (best control over cleanup)
    2) docx2pdf (Windows Word automation)
    3) LibreOffice (soffice) headless conversion
    """

    out_pdf_path.parent.mkdir(parents=True, exist_ok=True)

    # 1) Word automation via pywin32
    try:
        import win32com.client  # type: ignore

        wdFormatPDF = 17
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = None
        try:
            doc = word.Documents.Open(str(docx_path), ReadOnly=1)
            doc.SaveAs(str(out_pdf_path), FileFormat=wdFormatPDF)
        finally:
            try:
                if doc is not None:
                    doc.Close(False)
            finally:
                word.Quit()
        if out_pdf_path.exists():
            return
    except Exception:
        pass

    # 2) docx2pdf
    try:
        from docx2pdf import convert  # type: ignore

        # docx2pdf writes to the target directory using the docx stem.
        tmp_dir = out_pdf_path.parent
        convert(str(docx_path), str(tmp_dir))

        produced_pdf = tmp_dir / f"{docx_path.stem}.pdf"
        if not produced_pdf.exists():
            raise RuntimeError("docx2pdf did not produce the expected PDF")

        if produced_pdf.resolve() != out_pdf_path.resolve():
            if out_pdf_path.exists():
                out_pdf_path.unlink()
            produced_pdf.replace(out_pdf_path)

        return
    except Exception:
        pass

    soffice = shutil.which("soffice")
    if not soffice:
        raise RuntimeError(
            "Unable to convert DOCX to PDF (docx2pdf failed, and LibreOffice 'soffice' not found)."
        )

    out_pdf_path.parent.mkdir(parents=True, exist_ok=True)

    # LibreOffice writes <stem>.pdf into --outdir.
    subprocess.run(
        [
            soffice,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--convert-to",
            "pdf",
            "--outdir",
            str(out_pdf_path.parent),
            str(docx_path),
        ],
        check=True,
    )

    produced_pdf = out_pdf_path.parent / f"{docx_path.stem}.pdf"
    if not produced_pdf.exists():
        raise RuntimeError("LibreOffice did not produce the expected PDF")

    if produced_pdf.resolve() != out_pdf_path.resolve():
        if out_pdf_path.exists():
            out_pdf_path.unlink()
        produced_pdf.replace(out_pdf_path)


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    docs_dir = repo_root / "docs"

    in_md = docs_dir / "PROJECT_REPORT.md"
    # Keep final output as PDF only. We create a temporary DOCX solely for conversion.
    out_pdf = docs_dir / "ABL Activity Sample Report Updated.docx.pdf"

    meta = ReportMeta(
        title="ACTIVITY BASED LEARNING (ABL)",
        subtitle="ACTIVITY REPORT",
        project_name="Hospital OPD Management System",
        report_date=date.today().strftime("%d %b %Y"),
        student_name="<Your Name>",
        usn_or_id="<USN / ID>",
        class_and_section="<Class / Section>",
        department="<Department>",
        institution="<Institution Name>",
        guide_name="<Guide / Mentor Name>",
    )

    with tempfile.TemporaryDirectory(prefix="abl_report_", ignore_cleanup_errors=True) as tmp:
        tmp_docx = Path(tmp) / "ABL Activity Sample Report Updated.docx"
        build_docx(in_md, tmp_docx, meta)

        try:
            _convert_docx_to_pdf(tmp_docx, out_pdf)
            print(f"Created: {out_pdf}")
        except Exception as exc:
            print(f"PDF conversion skipped/failed: {exc}")


if __name__ == "__main__":
    main()
